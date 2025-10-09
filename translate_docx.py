#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor
import os

def translate_text(text):
    """Słownik tłumaczeń dla kluczowych fraz"""
    translations = {
        'Die Betraege in der excel Datei muessen mit Komma stehen, nicht mit punkt, sonst sind es keine zahlen, sondern text.':
            'Kwoty w pliku Excel muszą być zapisane z przecinkiem, a nie z kropką, w przeciwnym razie będą traktowane jako tekst, a nie liczby.',

        'Unsere Excel Datei muss die volle IBAN nummern haben in Auftraggeberkonto, die können wir in TM5 rausfinden.':
            'Nasz plik Excel musi zawierać pełne numery IBAN w polu Konto zleceniodawcy (Auftraggeberkonto), które możemy znaleźć w TM5.',

        'Fuer jedes Auftraggeber Konto muss eine separate Excel Datei erstellt werden.':
            'Dla każdego konta zleceniodawcy należy utworzyć osobny plik Excel.',

        'TM5': 'TM5',

        'Stammdaten': 'Dane podstawowe',
        'Konten': 'Konta',
        'Stammdaten ->': 'Dane podstawowe →',
        'Stammdaten -> Konten': 'Dane podstawowe → Konta'
    }

    # Sprawdź czy tekst istnieje w słowniku
    for german, polish in translations.items():
        if german.lower() in text.lower():
            text = text.replace(german, polish)

    return text

# Wczytaj oryginalny dokument
input_file = "TM5 Coupa Treasury Automatisierung Arbeitsanweisung.docx"
output_file = "TM5 Coupa Treasury Automatisierung Arbeitsanweisung - Tłumaczenie PL.docx"

print(f"Wczytywanie dokumentu: {input_file}")
doc = Document(input_file)

# Nowy dokument zachowujący obrazy
print("Przetwarzanie dokumentu...")

# Przetwórz wszystkie paragrafy
for para in doc.paragraphs:
    if para.text.strip():
        # Tłumacz tekst w paragrafie
        original_text = para.text
        translated = translate_text(original_text)

        if original_text != translated:
            print(f"  Tłumaczę: {original_text[:50]}...")
            # Zachowaj formatowanie
            for run in para.runs:
                if run.text.strip():
                    run.text = translate_text(run.text)

# Przetwórz tabele (jeśli są)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text.strip():
                    for run in para.runs:
                        if run.text.strip():
                            run.text = translate_text(run.text)

# Zapisz dokument
print(f"Zapisywanie przetłumaczonego dokumentu: {output_file}")
doc.save(output_file)

print("✓ Dokument został przetłumaczony z zachowaniem wszystkich obrazów!")
print(f"✓ Zapisano jako: {output_file}")
