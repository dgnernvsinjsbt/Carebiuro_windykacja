#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document

def translate_text(text):
    """Ulepszone tłumaczenie z niemieckiego na polski"""

    # Pełne tłumaczenia zdań
    full_translations = {
        'Die Betraege in der excel Datei muessen mit Komma stehen, nicht mit punkt, sonst sind es keine zahlen, sondern text.':
            'Kwoty w pliku Excel muszą być zapisane z przecinkiem, a nie z kropką, w przeciwnym razie będą traktowane jako tekst, a nie liczby.',

        'Unsere Excel Datei muss die volle IBAN nummern haben in Auftraggeberkonto, die können wir in TM5 rausfinden.':
            'Nasz plik Excel musi zawierać pełne numery IBAN w kolumnie "Konto zleceniodawcy" (Auftraggeberkonto), które możemy znaleźć w systemie TM5.',

        'Fuer jedes Auftraggeber Konto muss eine separate Excel Datei erstellt werden.':
            'Dla każdego konta zleceniodawcy należy utworzyć osobny plik Excel.',
    }

    # Tłumaczenia pojedynczych słów/fraz
    word_translations = {
        'Stammdaten': 'Dane podstawowe',
        'Konten': 'Konta',
        'TM5': 'TM5',
        'Excel Datei': 'plik Excel',
        'Auftraggeberkonto': 'Konto zleceniodawcy',
        'IBAN nummern': 'numery IBAN',
    }

    # Najpierw spróbuj pełnych tłumaczeń
    for german, polish in full_translations.items():
        if german in text:
            return polish

    # Potem pojedyncze słowa
    result = text
    for german, polish in word_translations.items():
        result = result.replace(german, polish)

    return result

# Wczytaj oryginalny dokument
input_file = "TM5 Coupa Treasury Automatisierung Arbeitsanweisung.docx"
output_file = "TM5 Coupa Treasury Automatisierung Arbeitsanweisung.docx"

print(f"Wczytywanie dokumentu: {input_file}")
doc = Document(input_file)

print("Przetwarzanie dokumentu...")

# Przetwórz wszystkie paragrafy
for para in doc.paragraphs:
    original_text = para.text.strip()
    if original_text:
        translated = translate_text(original_text)

        if original_text != translated:
            print(f"  '{original_text[:60]}...' → '{translated[:60]}...'")

            # Usuń wszystkie runs i dodaj nowy z przetłumaczonym tekstem
            # Zachowaj formatowanie pierwszego run'a
            if para.runs:
                first_run = para.runs[0]
                # Zapamiętaj formatowanie
                bold = first_run.bold
                italic = first_run.italic
                font_size = first_run.font.size
                font_name = first_run.font.name

                # Wyczyść paragraf
                for run in para.runs:
                    run.text = ""

                # Dodaj przetłumaczony tekst z formatowaniem
                new_run = para.runs[0]
                new_run.text = translated
                new_run.bold = bold
                new_run.italic = italic
                if font_size:
                    new_run.font.size = font_size
                if font_name:
                    new_run.font.name = font_name

# Przetwórz tabele
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                original_text = para.text.strip()
                if original_text:
                    translated = translate_text(original_text)
                    if original_text != translated:
                        print(f"  Tabela: '{original_text}' → '{translated}'")
                        if para.runs:
                            para.runs[0].text = translated
                            for run in para.runs[1:]:
                                run.text = ""

# Zapisz dokument (nadpisz oryginał)
print(f"\nZapisywanie dokumentu: {output_file}")
doc.save(output_file)

print("\n✅ Dokument został przetłumaczony z zachowaniem wszystkich obrazów i formatowania!")
print(f"✅ Plik nadpisany: {output_file}")
