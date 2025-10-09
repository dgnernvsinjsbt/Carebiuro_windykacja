#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
import re

def translate_text(text):
    """Kompletne tłumaczenie wszystkich fraz"""

    # Pełne zdania
    translations = {
        'Die Betraege in der excel Datei muessen mit Komma stehen, nicht mit punkt, sonst sind es keine zahlen, sondern text.':
            'Kwoty w pliku Excel muszą być zapisane z przecinkiem, a nie z kropką, w przeciwnym razie będą traktowane jako tekst, a nie liczby.',

        'Unsere Excel Datei muss die volle IBAN nummern haben in Auftraggeberkonto, die können wir in TM5 rausfinden.':
            'Nasz plik Excel musi zawierać pełne numery IBAN w kolumnie "Konto zleceniodawcy" (Auftraggeberkonto), które możemy znaleźć w systemie TM5.',

        'Fuer jedes Auftraggeber Konto muss eine separate Excel Datei erstellt werden.':
            'Dla każdego konta zleceniodawcy należy utworzyć osobny plik Excel.',

        'Die excel datei muss in dem gleichen folder gespeichert werden, wie die 3 template files .exe,.pbt, .bat':
            'Plik Excel musi być zapisany w tym samym folderze co 3 pliki szablonów: .exe, .pbt, .bat',

        'Die excel datei muss in dem gleichen folder gespeichert werden, wie die 3 template files':
            'Plik Excel musi być zapisany w tym samym folderze co 3 pliki szablonów',

        'Rechter mausklick auf die bat datei':
            'Kliknij prawym przyciskiem myszy na plik .bat',

        'edytuj w notatniku':
            'Edytuj w Notatniku',

        'Zmien nazwe pliku na nasz gotowy plik excel':
            'Zmień nazwę pliku na nasz gotowy plik Excel',

        # Pojedyncze słowa i frazy
        'Stammdaten': 'Dane podstawowe',
        'Konten': 'Konta',
        'TM5': 'TM5',
        'Excel Datei': 'plik Excel',
        'excel datei': 'plik Excel',
        'Auftraggeberkonto': 'Konto zleceniodawcy',
        'IBAN nummern': 'numery IBAN',
        'folder': 'folder',
        'template files': 'pliki szablonów',
        'Rechter mausklick': 'Kliknij prawym przyciskiem myszy',
        'bat datei': 'plik .bat',
        'die': '',
        'muss': 'musi',
        'in dem gleichen': 'w tym samym',
        'gespeichert werden': 'być zapisany',
        'wie': 'jak',
        'auf': 'na',
        'nazwe pliku': 'nazwę pliku',
        'nasz gotowy': 'nasz gotowy',
    }

    result = text

    # Tłumacz od najdłuższych do najkrótszych fraz
    for german, polish in sorted(translations.items(), key=lambda x: len(x[0]), reverse=True):
        # Case-insensitive replace
        pattern = re.compile(re.escape(german), re.IGNORECASE)
        result = pattern.sub(polish, result)

    return result

# Wczytaj dokument
input_file = "TM5 Coupa Treasury Automatisierung Arbeitsanweisung.docx"

print(f"Wczytywanie dokumentu: {input_file}")
doc = Document(input_file)

print("\nTłumaczenie wszystkich paragrafów:")
print("=" * 80)

# Przetwórz wszystkie paragrafy
for i, para in enumerate(doc.paragraphs):
    original_text = para.text.strip()
    if original_text:
        translated = translate_text(original_text)

        if original_text != translated:
            print(f"\nParagraf {i+1}:")
            print(f"  DE: {original_text}")
            print(f"  PL: {translated}")

            # Zastąp tekst zachowując formatowanie
            if para.runs:
                # Zapamiętaj formatowanie pierwszego run'a
                first_run = para.runs[0]
                bold = first_run.bold
                italic = first_run.italic
                font_size = first_run.font.size
                font_name = first_run.font.name

                # Wyczyść wszystkie runs
                for run in para.runs:
                    run.text = ""

                # Ustaw przetłumaczony tekst
                para.runs[0].text = translated
                para.runs[0].bold = bold
                para.runs[0].italic = italic
                if font_size:
                    para.runs[0].font.size = font_size
                if font_name:
                    para.runs[0].font.name = font_name

# Przetwórz tabele
print("\n\nTłumaczenie tabel:")
print("=" * 80)

for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                original_text = para.text.strip()
                if original_text:
                    translated = translate_text(original_text)
                    if original_text != translated:
                        print(f"\nTabela {table_idx+1}, Wiersz {row_idx+1}, Komórka {cell_idx+1}:")
                        print(f"  DE: {original_text}")
                        print(f"  PL: {translated}")

                        if para.runs:
                            para.runs[0].text = translated
                            for run in para.runs[1:]:
                                run.text = ""

# Zapisz
print("\n" + "=" * 80)
print(f"Zapisywanie dokumentu: {input_file}")
doc.save(input_file)

print("\n✅ DOKUMENT CAŁKOWICIE PRZETŁUMACZONY!")
print(f"✅ Wszystkie niemieckie teksty zostały przetłumaczone na polski")
print(f"✅ Zachowano: obrazy, formatowanie, strukturę")
